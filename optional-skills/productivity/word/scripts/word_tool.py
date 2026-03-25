#!/usr/bin/env python3
"""
Word Document Tool for Hermes Agent
-------------------------------------
Read, create, edit, search, and convert Word documents (.docx).

Requires: pip install python-docx

Usage:
  python3 word_tool.py info      <file>
  python3 word_tool.py read      <file> [--limit N] [--json]
  python3 word_tool.py create    <file> --content "text" [--title "Title"]
  python3 word_tool.py edit      <file> --find "old" --replace "new" [--all]
  python3 word_tool.py search    <file> <query> [--limit N]
  python3 word_tool.py add       <file> --text "paragraph" [--heading N] [--style NAME]
  python3 word_tool.py to-md     <file> [--output out.md]
  python3 word_tool.py from-md   <md_file> <docx_file> [--title "Title"]
  python3 word_tool.py extract   <file> [--output out.txt]
"""

import argparse
import json
import os
import re
import sys
from typing import Any, Dict, List, Optional


def _check_docx():
    try:
        import docx  # noqa: F401
        return True
    except ImportError:
        sys.exit("python-docx not installed.\n  pip install python-docx\n")


def _load_doc(path: str):
    import docx
    if not os.path.exists(path):
        sys.exit(f"File not found: {path}")
    try:
        return docx.Document(path)
    except Exception as e:
        sys.exit(f"Failed to open {path}: {e}")


# ---------------------------------------------------------------------------
# Commands
# ---------------------------------------------------------------------------

def cmd_info(args):
    """Show document metadata: paragraphs, tables, sections, styles."""
    _check_docx()
    doc = _load_doc(args.file)

    # Count elements
    paragraphs = len(doc.paragraphs)
    tables = len(doc.tables)
    sections = len(doc.sections)
    images = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            images += 1

    # Collect used styles
    styles_used = set()
    for p in doc.paragraphs:
        if p.style and p.style.name:
            styles_used.add(p.style.name)

    # Core properties
    props = doc.core_properties
    meta = {}
    for attr in ("author", "title", "subject", "keywords", "created", "modified", "last_modified_by"):
        val = getattr(props, attr, None)
        if val is not None:
            meta[attr] = str(val)

    # Section info
    sec_info = []
    for i, sec in enumerate(doc.sections):
        sec_info.append({
            "section": i + 1,
            "width_inches": round(sec.page_width.inches, 2) if sec.page_width else None,
            "height_inches": round(sec.page_height.inches, 2) if sec.page_height else None,
            "orientation": "landscape" if sec.orientation else "portrait",
        })

    print(json.dumps({
        "file": args.file,
        "paragraphs": paragraphs,
        "tables": tables,
        "sections": sections,
        "images": images,
        "styles_used": sorted(styles_used),
        "metadata": meta,
        "page_layout": sec_info,
    }, indent=2, default=str))


def cmd_read(args):
    """Read document content as structured data."""
    _check_docx()
    doc = _load_doc(args.file)

    content = []
    count = 0

    for p in doc.paragraphs:
        if count >= args.limit:
            break
        text = p.text.strip()
        if not text and not args.include_empty:
            continue

        entry = {"type": "paragraph", "text": text}
        if p.style and p.style.name:
            entry["style"] = p.style.name
        if p.style and p.style.name and p.style.name.startswith("Heading"):
            entry["type"] = "heading"
            try:
                entry["level"] = int(p.style.name.replace("Heading ", "").replace("Heading", "1"))
            except ValueError:
                entry["level"] = 1

        # Check for bold/italic runs
        has_bold = any(r.bold for r in p.runs if r.bold)
        has_italic = any(r.italic for r in p.runs if r.italic)
        if has_bold:
            entry["bold"] = True
        if has_italic:
            entry["italic"] = True

        content.append(entry)
        count += 1

    # Read tables
    for i, table in enumerate(doc.tables):
        if count >= args.limit:
            break
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        content.append({
            "type": "table",
            "table_index": i,
            "rows": len(rows),
            "columns": len(rows[0]) if rows else 0,
            "data": rows,
        })
        count += 1

    print(json.dumps({
        "file": args.file,
        "elements": len(content),
        "content": content,
    }, indent=2, default=str))


def _build_from_json(doc, elements: list) -> dict:
    """Build document structure from a JSON element array.

    Supported element types:
      {"type": "heading", "text": "...", "level": 1}
      {"type": "paragraph", "text": "..."}
      {"type": "bullets", "items": ["a", "b", "c"]}
      {"type": "numbered", "items": ["first", "second"]}
      {"type": "table", "headers": ["A","B"], "rows": [["1","2"],["3","4"]]}
      {"type": "quote", "text": "..."}
      {"type": "pagebreak"}
    """
    counts = {"headings": 0, "paragraphs": 0, "tables": 0, "lists": 0}

    for elem in elements:
        if not isinstance(elem, dict):
            continue
        etype = elem.get("type", "paragraph")

        if etype == "heading":
            level = min(max(int(elem.get("level", 1)), 1), 9)
            doc.add_heading(str(elem.get("text", "")), level=level)
            counts["headings"] += 1

        elif etype == "paragraph":
            text = str(elem.get("text", ""))
            p = doc.add_paragraph()
            # Support inline bold/italic via markdown-like syntax
            _add_formatted_runs(p, text)
            counts["paragraphs"] += 1

        elif etype == "bullets":
            for item in elem.get("items", []):
                doc.add_paragraph(str(item), style="List Bullet")
                counts["lists"] += 1

        elif etype == "numbered":
            for item in elem.get("items", []):
                doc.add_paragraph(str(item), style="List Number")
                counts["lists"] += 1

        elif etype == "table":
            headers = elem.get("headers", [])
            rows = elem.get("rows", [])
            # Also support flat "data" array where first row is headers
            data = elem.get("data")
            if data and not headers and not rows:
                headers = data[0] if data else []
                rows = data[1:] if len(data) > 1 else []

            all_rows = [headers] + rows if headers else rows
            if all_rows:
                cols = max(len(r) for r in all_rows)
                table = doc.add_table(rows=len(all_rows), cols=cols, style="Table Grid")
                for ri, row_data in enumerate(all_rows):
                    for ci, val in enumerate(row_data):
                        if ci < cols:
                            table.rows[ri].cells[ci].text = str(val)
                # Bold header row
                if headers:
                    for cell in table.rows[0].cells:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.bold = True
                counts["tables"] += 1

        elif etype == "quote":
            p = doc.add_paragraph(str(elem.get("text", "")))
            try:
                p.style = doc.styles["Quote"]
            except KeyError:
                pass  # Quote style not available in template
            counts["paragraphs"] += 1

        elif etype == "pagebreak":
            doc.add_page_break()

    return counts


def cmd_create(args):
    """Create a new Word document from text, JSON structure, or both."""
    _check_docx()
    import docx

    if os.path.exists(args.file) and not args.force:
        sys.exit(f"File already exists: {args.file}. Use --force to overwrite.")

    doc = docx.Document()

    if args.title:
        doc.add_heading(args.title, level=0)

    # JSON structure mode — build full document from elements
    if args.from_json:
        try:
            elements = json.loads(args.from_json)
            if not isinstance(elements, list):
                sys.exit("--from-json must be a JSON array of elements")
        except json.JSONDecodeError as e:
            sys.exit(f"Invalid JSON in --from-json: {e}")

        counts = _build_from_json(doc, elements)
        doc.save(args.file)
        print(json.dumps({
            "created": args.file,
            "title": args.title or None,
            "elements": counts,
            "total_paragraphs": len(doc.paragraphs),
        }, indent=2))
        return

    # Simple text mode
    if args.content:
        for para in args.content.split("\\n"):
            doc.add_paragraph(para)

    doc.save(args.file)
    print(json.dumps({
        "created": args.file,
        "title": args.title or None,
        "paragraphs": len(doc.paragraphs),
    }, indent=2))


def cmd_edit(args):
    """Find and replace text in a document."""
    _check_docx()
    doc = _load_doc(args.file)

    find_text = args.find
    replace_text = args.replace
    replacements = 0

    for p in doc.paragraphs:
        if find_text in p.text:
            # Replace in runs to preserve formatting
            for run in p.runs:
                if find_text in run.text:
                    if args.all:
                        count = run.text.count(find_text)
                        run.text = run.text.replace(find_text, replace_text)
                        replacements += count
                    else:
                        run.text = run.text.replace(find_text, replace_text, 1)
                        replacements += 1
                        if not args.all:
                            break
            if not args.all and replacements > 0:
                break

    # Also check table cells
    if args.all or replacements == 0:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if find_text in p.text:
                            for run in p.runs:
                                if find_text in run.text:
                                    count = run.text.count(find_text)
                                    run.text = run.text.replace(find_text, replace_text)
                                    replacements += count

    if replacements > 0:
        doc.save(args.file)

    print(json.dumps({
        "file": args.file,
        "find": find_text,
        "replace": replace_text,
        "replacements": replacements,
    }, indent=2))


def cmd_search(args):
    """Search for text in the document."""
    _check_docx()
    doc = _load_doc(args.file)

    query = args.query.lower()
    matches = []

    for i, p in enumerate(doc.paragraphs):
        if len(matches) >= args.limit:
            break
        if query in p.text.lower():
            matches.append({
                "paragraph": i + 1,
                "style": p.style.name if p.style else None,
                "text": p.text[:200] + ("..." if len(p.text) > 200 else ""),
            })

    # Search tables
    for ti, table in enumerate(doc.tables):
        if len(matches) >= args.limit:
            break
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                if query in cell.text.lower():
                    matches.append({
                        "table": ti + 1,
                        "row": ri + 1,
                        "column": ci + 1,
                        "text": cell.text[:200],
                    })
                    if len(matches) >= args.limit:
                        break

    print(json.dumps({
        "file": args.file,
        "query": args.query,
        "matches": len(matches),
        "results": matches,
    }, indent=2, default=str))


def cmd_add(args):
    """Add a paragraph or heading to the document."""
    _check_docx()
    doc = _load_doc(args.file)

    if args.heading is not None:
        doc.add_heading(args.text, level=args.heading)
        element_type = f"heading_{args.heading}"
    elif args.style:
        doc.add_paragraph(args.text, style=args.style)
        element_type = f"paragraph ({args.style})"
    else:
        doc.add_paragraph(args.text)
        element_type = "paragraph"

    doc.save(args.file)
    print(json.dumps({
        "file": args.file,
        "added": element_type,
        "text": args.text[:100],
        "total_paragraphs": len(doc.paragraphs),
    }, indent=2))


def cmd_to_md(args):
    """Convert Word document to Markdown."""
    _check_docx()
    doc = _load_doc(args.file)

    lines = []
    max_paragraphs = 10000  # safety cap

    for pi, p in enumerate(doc.paragraphs):
        if pi >= max_paragraphs:
            lines.append(f"\n[Truncated: {len(doc.paragraphs) - max_paragraphs} more paragraphs]")
            break
        text = p.text.strip()
        if not text:
            lines.append("")
            continue

        style = p.style.name if p.style else ""

        # Headings
        if style.startswith("Heading"):
            try:
                level = int(style.replace("Heading ", "").replace("Heading", "1"))
            except ValueError:
                level = 1
            lines.append(f"{'#' * level} {text}")
            continue

        # List items
        if style == "List Bullet" or style.startswith("List Bullet"):
            lines.append(f"- {text}")
            continue
        if style == "List Number" or style.startswith("List Number"):
            lines.append(f"1. {text}")
            continue

        # Blockquote
        if style == "Quote" or style == "Intense Quote":
            lines.append(f"> {text}")
            continue

        # Inline formatting via runs
        parts = []
        for run in p.runs:
            t = run.text
            if not t:
                continue
            if run.bold and run.italic:
                parts.append(f"***{t}***")
            elif run.bold:
                parts.append(f"**{t}**")
            elif run.italic:
                parts.append(f"*{t}*")
            else:
                parts.append(t)
        if parts:
            lines.append("".join(parts))
        else:
            lines.append(text)

    # Tables
    for table in doc.tables:
        lines.append("")
        for ri, row in enumerate(table.rows):
            cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
            lines.append("| " + " | ".join(cells) + " |")
            if ri == 0:
                lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
        lines.append("")

    md_content = "\n".join(lines)

    output = args.output
    if not output:
        output = os.path.splitext(args.file)[0] + ".md"

    with open(output, "w", encoding="utf-8") as f:
        f.write(md_content)

    print(json.dumps({
        "exported": output,
        "source": args.file,
        "lines": len(lines),
    }, indent=2))


def cmd_from_md(args):
    """Convert Markdown file to Word document."""
    _check_docx()
    import docx

    if not os.path.exists(args.md_file):
        sys.exit(f"Markdown file not found: {args.md_file}")
    if os.path.exists(args.docx_file):
        sys.exit(f"Output file already exists: {args.docx_file}. Remove it first.")

    with open(args.md_file, "r", encoding="utf-8") as f:
        md_content = f.read()

    doc = docx.Document()

    if args.title:
        doc.add_heading(args.title, level=0)

    lines = md_content.split("\n")
    in_table = False
    table_rows = []

    for line in lines:
        stripped = line.strip()

        # Skip table separator rows
        if re.match(r"^\|[\s\-|]+\|$", stripped):
            continue

        # Table row
        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(cells)
            continue
        elif in_table:
            # End of table — flush
            if table_rows:
                _add_table(doc, table_rows)
            in_table = False
            table_rows = []

        # Headings
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            doc.add_heading(heading_match.group(2), level=min(level, 9))
            continue

        # Horizontal rule
        if re.match(r"^[-*_]{3,}$", stripped):
            doc.add_paragraph("")
            continue

        # Blockquote
        if stripped.startswith("> "):
            p = doc.add_paragraph(stripped[2:])
            p.style = doc.styles["Quote"] if "Quote" in [s.name for s in doc.styles] else p.style
            continue

        # Bullet list
        if re.match(r"^[-*]\s+", stripped):
            text = re.sub(r"^[-*]\s+", "", stripped)
            doc.add_paragraph(text, style="List Bullet")
            continue

        # Numbered list
        if re.match(r"^\d+\.\s+", stripped):
            text = re.sub(r"^\d+\.\s+", "", stripped)
            doc.add_paragraph(text, style="List Number")
            continue

        # Empty line
        if not stripped:
            continue

        # Normal paragraph with inline formatting
        p = doc.add_paragraph()
        _add_formatted_runs(p, stripped)

    # Flush any remaining table
    if in_table and table_rows:
        _add_table(doc, table_rows)

    doc.save(args.docx_file)
    print(json.dumps({
        "created": args.docx_file,
        "source": args.md_file,
        "paragraphs": len(doc.paragraphs),
    }, indent=2))


def _add_table(doc, rows):
    """Add a table to the document from row data."""
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols, style="Table Grid")
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            if ci < cols:
                table.rows[ri].cells[ci].text = val


def _add_formatted_runs(paragraph, text):
    """Parse inline Markdown formatting and add styled runs."""
    # Split on bold/italic markers
    parts = re.split(r"(\*{1,3}[^*]+\*{1,3})", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("***") and part.endswith("***"):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def cmd_extract(args):
    """Extract plain text from a Word document."""
    _check_docx()
    doc = _load_doc(args.file)

    texts = []
    max_extract = 10000  # safety cap
    for i, p in enumerate(doc.paragraphs):
        if i >= max_extract:
            texts.append(f"[Truncated: {len(doc.paragraphs) - max_extract} more paragraphs]")
            break
        texts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            texts.append("\t".join(cell.text for cell in row.cells))

    full_text = "\n".join(texts)

    output = args.output
    if output:
        with open(output, "w", encoding="utf-8") as f:
            f.write(full_text)
        print(json.dumps({
            "extracted": output,
            "source": args.file,
            "characters": len(full_text),
        }, indent=2))
    else:
        print(json.dumps({
            "file": args.file,
            "characters": len(full_text),
            "text": full_text[:5000] + ("..." if len(full_text) > 5000 else ""),
        }, indent=2))


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        prog="word_tool.py",
        description="Word document tool — read, create, edit, search, convert .docx files",
    )
    sub = parser.add_subparsers(dest="command", required=True)

    p_info = sub.add_parser("info", help="Document metadata: paragraphs, tables, styles, images")
    p_info.add_argument("file")

    p_read = sub.add_parser("read", help="Read document content as structured data")
    p_read.add_argument("file")
    p_read.add_argument("--limit", type=int, default=200, help="Max elements (default: 200)")
    p_read.add_argument("--include-empty", action="store_true", help="Include empty paragraphs")

    p_create = sub.add_parser("create", help="Create a new Word document")
    p_create.add_argument("file")
    p_create.add_argument("--content", help="Text content (use \\n for line breaks)")
    p_create.add_argument("--title", help="Document title (added as Heading 0)")
    p_create.add_argument("--from-json", help="JSON array of elements: headings, paragraphs, bullets, tables, etc.")
    p_create.add_argument("--force", action="store_true", help="Overwrite existing file")

    p_edit = sub.add_parser("edit", help="Find and replace text")
    p_edit.add_argument("file")
    p_edit.add_argument("--find", required=True, help="Text to find")
    p_edit.add_argument("--replace", required=True, help="Replacement text")
    p_edit.add_argument("--all", action="store_true", help="Replace all occurrences")

    p_search = sub.add_parser("search", help="Search for text in the document")
    p_search.add_argument("file")
    p_search.add_argument("query")
    p_search.add_argument("--limit", type=int, default=50, help="Max results (default: 50)")

    p_add = sub.add_parser("add", help="Add a paragraph or heading")
    p_add.add_argument("file")
    p_add.add_argument("--text", required=True, help="Text to add")
    p_add.add_argument("--heading", type=int, help="Add as heading (level 1-9)")
    p_add.add_argument("--style", help="Paragraph style name")

    p_tomd = sub.add_parser("to-md", help="Convert Word to Markdown")
    p_tomd.add_argument("file")
    p_tomd.add_argument("--output", help="Output .md path (default: auto)")

    p_frommd = sub.add_parser("from-md", help="Convert Markdown to Word")
    p_frommd.add_argument("md_file")
    p_frommd.add_argument("docx_file")
    p_frommd.add_argument("--title", help="Document title")

    p_extract = sub.add_parser("extract", help="Extract plain text from document")
    p_extract.add_argument("file")
    p_extract.add_argument("--output", help="Output .txt path (default: print to stdout)")

    args = parser.parse_args()
    dispatch = {
        "info": cmd_info, "read": cmd_read, "create": cmd_create,
        "edit": cmd_edit, "search": cmd_search, "add": cmd_add,
        "to-md": cmd_to_md, "from-md": cmd_from_md, "extract": cmd_extract,
    }
    dispatch[args.command](args)


if __name__ == "__main__":
    main()
